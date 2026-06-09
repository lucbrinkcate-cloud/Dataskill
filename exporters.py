from __future__ import annotations

import html
import math
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence


class HTMLTextExtractor(HTMLParser):
    """Small dependency-free HTML to structured text extractor.

    This is not a full HTML renderer. It creates readable text for DOCX/PDF fallback
    exports while keeping the self-contained HTML file as the source of truth.
    """

    BLOCK_TAGS = {"p", "div", "section", "article", "header", "footer", "main", "br", "tr", "table", "ul", "ol"}
    HEADING_TAGS = {"h1", "h2", "h3", "h4"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: List[Dict[str, str]] = []
        self.current: List[str] = []
        self.current_tag: str = "p"
        self.in_script_style = False
        self.last_was_break = True

    def _flush(self) -> None:
        text = " ".join("".join(self.current).split())
        if text:
            self.parts.append({"tag": self.current_tag or "p", "text": html.unescape(text)})
        self.current = []
        self.current_tag = "p"

    def handle_starttag(self, tag: str, attrs: List[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style"}:
            self.in_script_style = True
            return
        if tag in self.HEADING_TAGS:
            self._flush()
            self.current_tag = tag
        elif tag == "li":
            self._flush()
            self.current_tag = "li"
            self.current.append("• ")
        elif tag in {"td", "th"}:
            if self.current and not self.current[-1].endswith(" | "):
                self.current.append(" | ")
        elif tag in self.BLOCK_TAGS and self.current:
            self._flush()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style"}:
            self.in_script_style = False
            return
        if tag in self.HEADING_TAGS or tag in {"p", "div", "section", "tr", "li"}:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self.in_script_style:
            return
        if data and data.strip():
            self.current.append(data)

    def close(self) -> None:
        super().close()
        self._flush()


def extract_html_parts(html_text: str) -> List[Dict[str, str]]:
    parser = HTMLTextExtractor()
    parser.feed(html_text)
    parser.close()
    # Remove duplicate adjacent lines created by nested div/table parsing.
    out: List[Dict[str, str]] = []
    prev = ""
    for part in parser.parts:
        text = part["text"].strip()
        if not text or text == prev:
            continue
        out.append({"tag": part.get("tag", "p"), "text": text})
        prev = text
    return out


def html_to_text(html_text: str, max_lines: int = 5000) -> str:
    return "\n".join(part["text"] for part in extract_html_parts(html_text)[:max_lines])


def html_file_to_docx(html_path: str, docx_path: str) -> str:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as exc:
        raise RuntimeError("DOCX export requires python-docx. Install requirements.txt first.") from exc

    html_text = Path(html_path).read_text(encoding="utf-8", errors="ignore")
    parts = extract_html_parts(html_text)
    doc = Document()
    styles = doc.styles
    try:
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10)
    except Exception:
        pass
    for part in parts:
        tag = part.get("tag", "p")
        text = part.get("text", "").strip()
        if not text:
            continue
        if tag == "h1":
            doc.add_heading(text, level=1)
        elif tag == "h2":
            doc.add_heading(text, level=2)
        elif tag == "h3":
            doc.add_heading(text, level=3)
        elif tag == "li" or text.startswith("• "):
            doc.add_paragraph(text[2:] if text.startswith("• ") else text, style="List Bullet")
        elif " | " in text and len(text) < 1000:
            # Simple readable table-row fallback.
            doc.add_paragraph(text)
        else:
            doc.add_paragraph(text)
    out = Path(docx_path).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


def html_file_to_pdf(html_path: str, pdf_path: str) -> str:
    """Export HTML to PDF.

    Preferred: WeasyPrint if installed. Fallback: ReportLab text PDF if installed.
    Last-resort: a minimal dependency-free text PDF writer.
    """
    html_text = Path(html_path).read_text(encoding="utf-8", errors="ignore")
    out = Path(pdf_path).expanduser().resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    try:
        from weasyprint import HTML  # type: ignore
        HTML(filename=str(Path(html_path).resolve())).write_pdf(str(out))
        return str(out)
    except Exception:
        pass

    text = html_to_text(html_text, max_lines=4000)
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore
        from xml.sax.saxutils import escape as xml_escape

        doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=40, leftMargin=40, topMargin=42, bottomMargin=42)
        styles = getSampleStyleSheet()
        story = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            story.append(Paragraph(xml_escape(line), styles["Normal"]))
            story.append(Spacer(1, 4))
        doc.build(story)
        return str(out)
    except Exception:
        _write_minimal_text_pdf(text, out)
        return str(out)


def _pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_minimal_text_pdf(text: str, out: Path) -> None:
    # Minimal multipage A4 PDF with Helvetica text. Good enough as fallback.
    lines = []
    for raw in text.splitlines():
        raw = raw.strip()
        if not raw:
            lines.append("")
            continue
        while len(raw) > 95:
            lines.append(raw[:95])
            raw = raw[95:]
        lines.append(raw)
    lines_per_page = 52
    pages = [lines[i : i + lines_per_page] for i in range(0, len(lines), lines_per_page)] or [["No content"]]

    objects: List[bytes] = []
    # 1 catalog, 2 pages, 3 font, then page/content pairs.
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{4 + i*2} 0 R" for i in range(len(pages)))
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for i, page_lines in enumerate(pages):
        page_obj_num = 4 + i * 2
        content_obj_num = page_obj_num + 1
        objects.append(f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 3 0 R >> >> /Contents {content_obj_num} 0 R >>".encode())
        commands = ["BT", "/F1 10 Tf", "50 800 Td", "14 TL"]
        for line in page_lines:
            commands.append(f"({_pdf_escape(line)}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands).encode("latin-1", errors="replace")
        objects.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for idx, obj in enumerate(objects, 1):
        offsets.append(len(pdf))
        pdf.extend(f"{idx} 0 obj\n".encode())
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_pos = len(pdf)
    pdf.extend(f"xref\n0 {len(objects)+1}\n".encode())
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode())
    pdf.extend(f"trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode())
    out.write_bytes(bytes(pdf))


def export_html(html_path: str, formats: Sequence[str]) -> Dict[str, Any]:
    """Create requested export formats from an HTML file.

    Returns a dict with ok flag and generated path list.
    """
    base = Path(html_path).expanduser().resolve()
    fmt_set = {f.lower().strip() for f in formats if f and f.strip()}
    if not fmt_set:
        fmt_set = {"html"}
    outputs: Dict[str, str] = {}
    errors: Dict[str, str] = {}
    if "html" in fmt_set:
        outputs["html"] = str(base)
    if "pdf" in fmt_set:
        try:
            outputs["pdf"] = html_file_to_pdf(str(base), str(base.with_suffix(".pdf")))
        except Exception as exc:
            errors["pdf"] = str(exc)
    if "docx" in fmt_set or "word" in fmt_set:
        try:
            outputs["docx"] = html_file_to_docx(str(base), str(base.with_suffix(".docx")))
        except Exception as exc:
            errors["docx"] = str(exc)
    return {"ok": not errors, "outputs": outputs, "errors": errors}
